import os
import time
import json
from pathlib import Path
from anthropic import Anthropic
from loguru import logger
from dotenv import load_dotenv
from utils.models import Document, AuditLog, DocumentStatus, InferenceMode, get_session

load_dotenv()

client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
CONFIDENCE_THRESHOLD = float(os.getenv("CONFIDENCE_THRESHOLD", 0.75))

# expanded categories — resume was falling through to unknown before
DOC_CATEGORIES = [
    "financial_statement", "insurance_claim", "loan_application",
    "audit_evidence", "healthcare_record", "legal_contract",
    "invoice", "resume_cv", "report", "email", "unknown"
]

EXTRACT_TOOLS = [{
    "name": "classify_document",
    "description": "Classify doc type and extract metadata",
    "input_schema": {
        "type": "object",
        "properties": {
            "category": {
                "type": "string",
                "enum": DOC_CATEGORIES,
                "description": "Document category. Use resume_cv for resumes, CVs, job applications."
            },
            "confidence": {"type": "number", "description": "0.0-1.0 confidence in classification"},
            "key_entities": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string"},
            "risk_indicators": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Only flag genuine anomalies relevant to the document type. Do not flag resumes as high risk."
            },
            "anomaly_detected": {"type": "boolean"}
        },
        "required": ["category", "confidence", "key_entities", "summary", "anomaly_detected"]
    }
}]


def read_file(path: str) -> tuple[str, str]:
    suffix = Path(path).suffix.lower()

    if suffix == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        return text.strip(), "pdf"

    if suffix == ".docx":
        import docx as _docx
        doc = _docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()), "docx"

    if suffix == ".txt":
        return open(path, encoding="utf-8").read().strip(), "txt"

    raise ValueError(f"can't handle {suffix}")


def classify(text: str, filename: str) -> dict:
    # 8k chars is enough for most docs
    snippet = text[:8000] + ("..." if len(text) > 8000 else "")

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        system="""Enterprise document classifier. Use classify_document tool.
Category guidance:
- resume_cv: any resume, CV, job application, professional profile
- invoice: bills, purchase orders, payment requests
- financial_statement: balance sheets, income statements, financial reports
- insurance_claim: insurance forms, claim submissions
- loan_application: loan forms, mortgage applications, credit applications
- audit_evidence: audit reports, compliance documents
- healthcare_record: medical records, patient documents, clinical notes
- legal_contract: contracts, agreements, legal documents
- report: analytical reports, research, summaries
- email: email correspondence

Risk indicators: only flag genuine anomalies relevant to the document type.
Resumes should almost never have risk indicators.""",
        tools=EXTRACT_TOOLS,
        messages=[{"role": "user", "content": f"Classify and extract from this document.\n\nFile: {filename}\n\n{snippet}"}]
    )

    for block in resp.content:
        if block.type == "tool_use" and block.name == "classify_document":
            return block.input

    # TODO: add retry
    logger.warning(f"Claude skipped tool call on {filename}, using fallback")
    return {
        "category": "unknown", "confidence": 0.5,
        "key_entities": [], "summary": "classification failed",
        "risk_indicators": [], "anomaly_detected": False
    }


def process_document(file_path: str, inference_mode: str = "realtime") -> dict:
    t0 = time.time()
    session = get_session()
    fname = Path(file_path).name

    try:
        raw, ftype = read_file(file_path)
        output = classify(raw, fname)
        latency = (time.time() - t0) * 1000
        conf = output.get("confidence", 0)

        status = DocumentStatus.REVIEW_REQUIRED if conf < CONFIDENCE_THRESHOLD else DocumentStatus.COMPLETED

        doc = Document(
            filename=fname,
            file_type=ftype,
            document_category=output.get("category"),
            raw_text=raw,
            structured_output=output,
            status=status,
            inference_mode=InferenceMode(inference_mode),
            processing_latency_ms=latency
        )
        session.add(doc)
        session.add(AuditLog(
            document_id=doc.id,
            action="processed",
            details={"cat": output.get("category"), "conf": conf, "ms": latency}
        ))
        session.commit()

        logger.info(f"{fname} → {output.get('category')} ({conf:.2f}) {latency:.0f}ms")
        return {
            "document_id": doc.id, "filename": fname,
            "category": output.get("category"), "confidence": conf,
            "structured_output": output, "status": status,
            "latency_ms": latency, "flagged_for_review": status == DocumentStatus.REVIEW_REQUIRED
        }

    except Exception as e:
        session.rollback()
        logger.error(f"failed on {fname}: {e}")
        doc = Document(filename=fname, file_type="unknown", status=DocumentStatus.FAILED)
        session.add(doc)
        session.commit()
        return {"document_id": doc.id, "filename": fname, "status": "failed", "error": str(e)}

    finally:
        session.close()


def process_batch(paths: list[str]) -> list[dict]:
    logger.info(f"batch: {len(paths)} docs")
    return [process_document(p, "batch") for p in paths]
